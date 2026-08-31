"""
Generate BPS Standard DOCX Documents for Phase 4: System Testing & Verification
Sistem Monitoring SE2026 BPS Kabupaten Penajam Paser Utara ("Pananyo Taka")
Format: Hitam-putih, tanpa fill/background, hanya garis tabel standar.
"""

import os
import re
import sys
import shutil
import docx

if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass

from docx.shared import Inches, Pt, Emu, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Colour constants (hitam-putih saja) ──────────────────────────────────────
BLACK  = RGBColor(0, 0, 0)
WHITE  = RGBColor(255, 255, 255)

# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcMar)

def _set_table_grid_borders(table, color='000000', sz='4'):
    """Full-grid black borders on all sides."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)

def _para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, sp_before=1, sp_after=2, ls=1.15):
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(sp_before)
    pf.space_after  = Pt(sp_after)
    pf.line_spacing = ls

def _run(para, text, bold=False, italic=False, size=Pt(10), underline=False):
    r = para.add_run(text)
    r.font.name    = 'Calibri'
    r.font.size    = size
    r.font.color.rgb = BLACK
    r.bold         = bold
    r.italic       = italic
    r.underline    = underline
    return r

# ── Document factory ──────────────────────────────────────────────────────────

def create_document():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        # Running header
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run('BPS Kabupaten Penajam Paser Utara | Dashboard Monitoring SE2026')
        r.font.name = 'Calibri'; r.font.size = Pt(8); r.font.color.rgb = BLACK
        # Footer
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run('Laporan Pengujian dan Verifikasi Sistem (SDLC Phase 4) — Pananyo Taka v1.0.0')
        r.font.name = 'Calibri'; r.font.size = Pt(8); r.font.color.rgb = BLACK
    return doc

# ── Header banner (teks saja, tanpa fill) ─────────────────────────────────────

def add_header_banner(doc, title, subtitle, subsubtitle=None):
    """Judul dokumen — teks centered saja, tanpa kotak berwarna."""
    # Institusi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_format(p, sp_before=0, sp_after=1)
    _run(p, 'BADAN PUSAT STATISTIK KABUPATEN PENAJAM PASER UTARA', bold=True, size=Pt(12))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_format(p, sp_before=0, sp_after=2)
    _run(p, 'PELATIHAN DASAR CPNS GOLONGAN III TAHUN 2026', size=Pt(10))

    # Garis tipis
    p_hr = doc.add_paragraph()
    _para_format(p_hr, sp_before=0, sp_after=0)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Judul dokumen
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_format(p, sp_before=6, sp_after=2)
    _run(p, title, bold=True, size=Pt(13))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_format(p, sp_before=0, sp_after=2)
    _run(p, subtitle, bold=True, size=Pt(10.5))

    if subsubtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_format(p, sp_before=0, sp_after=8)
        _run(p, subsubtitle, italic=True, size=Pt(9.5))

    # Garis bawah judul
    p_hr2 = doc.add_paragraph()
    _para_format(p_hr2, sp_before=0, sp_after=8)
    pPr2 = p_hr2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    bot2 = OxmlElement('w:bottom')
    bot2.set(qn('w:val'), 'single'); bot2.set(qn('w:sz'), '6')
    bot2.set(qn('w:space'), '1'); bot2.set(qn('w:color'), '000000')
    pBdr2.append(bot2)
    pPr2.append(pBdr2)

# ── Markdown → DOCX parser ────────────────────────────────────────────────────

def _clean(text):
    """Strip markdown bold/italic markers and backticks."""
    return re.sub(r'[`]', '', re.sub(r'\*+', '', text))

def parse_markdown_to_docx(doc, md_content):
    lines = md_content.split('\n')
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip metadata header lines at the top
        if i < 18 and (stripped.startswith('# ') or stripped.startswith('## ') or
                       stripped.startswith('### ') or stripped.startswith('---')):
            i += 1
            continue

        # Code blocks → monospace table
        if stripped.startswith('```'):
            if in_code:
                in_code = False
                code_text = '\n'.join(code_lines)
                tbl = doc.add_table(rows=1, cols=1)
                _set_table_grid_borders(tbl)
                c = tbl.cell(0, 0)
                _set_cell_margins(c)
                cp = c.paragraphs[0]
                _para_format(cp, sp_before=0, sp_after=0, ls=1.0)
                r = cp.add_run(code_text)
                r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = BLACK
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                code_lines = []
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in stripped and ('|---|' in stripped or
                (i + 1 < len(lines) and '|---|' in lines[i + 1])):
            table_lines = [stripped]
            i += 1
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1

            parsed_rows = []
            for tl in table_lines:
                if re.match(r'^[\|\s\-:]+$', tl):
                    continue
                cols = [c.strip() for c in tl.strip('|').split('|')]
                parsed_rows.append(cols)

            if parsed_rows:
                num_cols = max(len(r) for r in parsed_rows)
                tbl = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                _set_table_grid_borders(tbl)

                for r_idx, row_data in enumerate(parsed_rows):
                    is_header = (r_idx == 0)
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                            _set_cell_margins(cell, top=60, bottom=60, left=90, right=90)
                            p = cell.paragraphs[0]
                            _para_format(p, sp_before=0, sp_after=0, ls=1.12)
                            clean = _clean(cell_value)
                            # Detect success keywords for bold only (no colour)
                            is_success = any(w in clean for w in [
                                'PASSED', 'ACCEPT', 'LULUS', 'SOLVED',
                                'VERIFIED', '100%', 'TERVERIFIKASI', 'TERSELESAIKAN'])
                            r = p.add_run(clean)
                            r.font.name  = 'Calibri'
                            r.font.size  = Pt(8.5)
                            r.font.color.rgb = BLACK
                            r.font.bold  = (is_header or is_success)

                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            _para_format(p, sp_before=12, sp_after=3)
            p.paragraph_format.keep_with_next = True
            _run(p, stripped[2:].strip(), bold=True, size=Pt(12))
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            _para_format(p, sp_before=10, sp_after=3)
            p.paragraph_format.keep_with_next = True
            _run(p, stripped[3:].strip(), bold=True, size=Pt(11))
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            _para_format(p, sp_before=8, sp_after=2)
            p.paragraph_format.keep_with_next = True
            _run(p, stripped[4:].strip(), bold=True, size=Pt(10.5))
        elif stripped.startswith('> '):
            # Callout: simple indented box (garis kiri saja)
            callout = _clean(stripped[2:].strip())
            tbl = doc.add_table(rows=1, cols=1)
            _set_table_grid_borders(tbl)
            c = tbl.cell(0, 0)
            _set_cell_margins(c, top=70, bottom=70, left=100, right=100)
            cp = c.paragraphs[0]
            _para_format(cp, sp_before=0, sp_after=0, ls=1.15)
            _run(cp, callout, italic=True, size=Pt(9.5))
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _para_format(p, sp_before=1, sp_after=1, ls=1.12)
            parts = stripped[2:].strip().split('**')
            for idx, part in enumerate(parts):
                _run(p, _clean(part), bold=(idx % 2 == 1), size=Pt(9.5))
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph()
            _para_format(p, sp_before=1, sp_after=1, ls=1.12)
            m = re.match(r'^(\d+\.)\s*(.*)', stripped)
            if m:
                _run(p, m.group(1) + ' ', bold=True, size=Pt(9.5))
                parts = m.group(2).split('**')
                for idx, part in enumerate(parts):
                    _run(p, _clean(part), bold=(idx % 2 == 1), size=Pt(9.5))
        elif stripped.startswith('<table') or stripped.startswith('<tr') or \
             stripped.startswith('<td') or stripped.startswith('</') or \
             stripped.startswith('<br'):
            pass  # skip raw HTML signature tables in markdown
        elif stripped and not stripped.startswith('---'):
            p = doc.add_paragraph()
            _para_format(p, sp_before=1, sp_after=2, ls=1.15)
            parts = stripped.split('**')
            for idx, part in enumerate(parts):
                _run(p, _clean(part), bold=(idx % 2 == 1), size=Pt(9.5))

        i += 1

# ── Signature block (hitam-putih, tanpa fill) ─────────────────────────────────

def add_signature_block(doc):
    doc.add_paragraph().paragraph_format.space_before = Pt(14)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_grid_borders(tbl)
    tbl.columns[0].width = Inches(3.2)
    tbl.columns[1].width = Inches(3.2)

    def sig_cell(cell, heading, name_line, date_line=None):
        _set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_format(p, sp_before=0, sp_after=2, ls=1.15)
        if date_line:
            _run(p, date_line + '\n', size=Pt(9.5))
        _run(p, heading + '\n', bold=True, size=Pt(9.5))
        _run(p, 'Pranata Komputer Ahli Pertama\n', size=Pt(9))
        _run(p, 'BPS Kabupaten Penajam Paser Utara\n\n\n\n\n', size=Pt(9))
        _run(p, name_line, bold=True, underline=True, size=Pt(9.5))

    sig_cell(tbl.cell(0, 0),
             heading='Mentor / Pembimbing Aktualisasi',
             name_line='BAIHAQI ILHAM SYAH, S.Tr.Stat.')
    sig_cell(tbl.cell(0, 1),
             heading='Penyusun / Perekayasa Sistem',
             name_line='YAHYA ABDURROHMAN',
             date_line='Penajam Paser Utara, 25 Agustus 2026')

# ── Safe save (handles Word file-lock) ───────────────────────────────────────

def _safe_save(doc, target_path):
    """Save doc via temp file to avoid PermissionError when Word has target open."""
    import tempfile, shutil, time
    folder = os.path.dirname(target_path) or '.'
    os.makedirs(folder, exist_ok=True)
    fd, tmp = tempfile.mkstemp(suffix='.docx', dir=folder)
    os.close(fd)
    doc.save(tmp)
    for attempt in range(6):
        try:
            shutil.move(tmp, target_path)
            print(f'    [OK] Berhasil dibuat: {target_path}')
            return
        except (PermissionError, OSError):
            if attempt < 5:
                time.sleep(2)
    alt = target_path.replace('.docx', '_NEW.docx')
    shutil.move(tmp, alt)
    print(f'    [WARN] File terkunci Word, disimpan sebagai: {alt}')

# ── Main build ────────────────────────────────────────────────────────────────

def generate_all_docx():
    print('Memulai pembuatan seluruh dokumen DOCX Phase 4 (Format Hitam-Putih)...')

    base = os.path.join('laporan', 'OUTPUT_TAHAPAN_KEGIATAN',
                        'Kegiatan_4_System_Testing_and_Verification')

    docs = [
        {
            'md':   os.path.join(base, 'Tahapan_4.1_Pengujian_Internal_BlackBox_dan_Sinkronisasi',
                                 'Catatan_Hasil_Uji_Coba_Performa_BlackBox_dan_Sinkronisasi.md'),
            'docx': os.path.join(base, 'Tahapan_4.1_Pengujian_Internal_BlackBox_dan_Sinkronisasi',
                                 'Catatan_Hasil_Uji_Coba_Performa_BlackBox_dan_Sinkronisasi.docx'),
            'title':    'CATATAN HASIL UJI COBA PERFORMA SISTEM',
            'subtitle': 'Black-Box Functional Testing & Data Synchronization Pipelines',
            'sub3':     'Tahapan 4.1: Pengujian Internal Sistem (Pressman & Maxim, 2015)',
        },
        {
            'md':   os.path.join(base, 'Tahapan_4.2_Pengujian_Performa_LoadSpeed_dan_Lighthouse',
                                 'Laporan_Hasil_Pengujian_Performa_Dasbor_LoadSpeed_dan_Lighthouse.md'),
            'docx': os.path.join(base, 'Tahapan_4.2_Pengujian_Performa_LoadSpeed_dan_Lighthouse',
                                 'Laporan_Hasil_Pengujian_Performa_Dasbor_LoadSpeed_dan_Lighthouse.docx'),
            'title':    'LAPORAN HASIL PENGUJIAN PERFORMA DASBOR',
            'subtitle': 'Concurrency Stress Testing, Load Speed & Google Lighthouse Audit',
            'sub3':     'Tahapan 4.2: Pengujian Performa dan Responsivitas Sistem',
        },
        {
            'md':   os.path.join(base, 'Tahapan_4.3_Uji_Coba_Terbatas_UAT_dan_Bug_Fixing',
                                 'Berita_Acara_UAT_dan_Catatan_Perbaikan_Bug.md'),
            'docx': os.path.join(base, 'Tahapan_4.3_Uji_Coba_Terbatas_UAT_dan_Bug_Fixing',
                                 'Berita_Acara_UAT_dan_Catatan_Perbaikan_Bug.docx'),
            'title':    'BERITA ACARA UAT & CATATAN PERBAIKAN BUG',
            'subtitle': 'User Acceptance Testing (18 Skenario) & Defect Resolution Log',
            'sub3':     'Tahapan 4.3: Pelaksanaan Uji Coba Terbatas & Bug Fixing',
        },
        {
            'md':   os.path.join(base, 'Tahapan_4.4_Berita_Acara_Pengujian_dan_Lembar_Persetujuan',
                                 'Laporan_Akhir_Hasil_Pengujian_dan_Lembar_Persetujuan_Kelayakan.md'),
            'docx': os.path.join(base, 'Tahapan_4.4_Berita_Acara_Pengujian_dan_Lembar_Persetujuan',
                                 'Laporan_Akhir_Hasil_Pengujian_dan_Lembar_Persetujuan_Kelayakan.docx'),
            'title':    'LAPORAN AKHIR HASIL PENGUJIAN & VALIDASI KELAYAKAN',
            'subtitle': 'Evaluasi Standar Mutu ISO/IEC 25010 & Go-Live Readiness',
            'sub3':     'Tahapan 4.4: Berita Acara Pengujian & Lembar Persetujuan Kelayakan',
        },
        {
            'md':   os.path.join(base, 'Master_Laporan_Phase_4',
                                 'Laporan_Pengujian_dan_Verifikasi_Sistem_SE2026_PPU_Phase4.md'),
            'docx': os.path.join(base, 'Master_Laporan_Phase_4',
                                 'Laporan_Pengujian_dan_Verifikasi_Sistem_SE2026_PPU_Phase4.docx'),
            'title':    'LAPORAN PENGUJIAN DAN VERIFIKASI SISTEM (PHASE 4)',
            'subtitle': "Dashboard Monitoring Sensus Ekonomi 2026 'Pananyo Taka' v1.0.0",
            'sub3':     'SDLC Phase 4: System Testing & Verification (Pressman & Maxim, 2015)',
        },
    ]

    for item in docs:
        print(f"--> Memproses: {item['docx']}")
        doc = create_document()
        add_header_banner(doc, item['title'], item['subtitle'], item.get('sub3'))
        with open(item['md'], 'r', encoding='utf-8') as f:
            md = f.read()
        parse_markdown_to_docx(doc, md)
        add_signature_block(doc)
        _safe_save(doc, item['docx'])

    root_master = 'Laporan_Pengujian_dan_Verifikasi_Sistem_SE2026_PPU_Phase4.docx'
    src_master = os.path.join(base, 'Master_Laporan_Phase_4',
                              'Laporan_Pengujian_dan_Verifikasi_Sistem_SE2026_PPU_Phase4.docx')
    # Use _NEW fallback if master was locked
    if not os.path.exists(src_master):
        src_master = src_master.replace('.docx', '_NEW.docx')
    import time
    for attempt in range(6):
        try:
            shutil.copyfile(src_master, root_master)
            print(f'[OK] Master DOCX berhasil disalin ke root: {root_master}')
            break
        except (PermissionError, OSError):
            if attempt < 5:
                time.sleep(2)
            else:
                alt_root = root_master.replace('.docx', '_NEW.docx')
                shutil.copyfile(src_master, alt_root)
                print(f'[WARN] Root master terkunci, disimpan sebagai: {alt_root}')
    print('\nSeluruh 5 dokumen DOCX Phase 4 berhasil dibuat dengan sukses!')


if __name__ == '__main__':
    generate_all_docx()
