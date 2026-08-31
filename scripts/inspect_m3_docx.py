import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree

doc = Document(r'laporan\03_Laporan_Mingguan_Latsar\M3 - Laporan Mingguan Pelaksanaan Aktualisasi Pelatihan Dasar CPNS BPS Tahun 2026.docx')

print('=== SECTIONS ===')
for i, s in enumerate(doc.sections):
    print(f'  Section {i}: top={s.top_margin}, bottom={s.bottom_margin}, left={s.left_margin}, right={s.right_margin}')
    print(f'  Page width={s.page_width}, height={s.page_height}')

print()
print('=== FIRST 5 PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs[:5]):
    print(f'Para {i}: align={para.alignment}, text="{para.text[:80]}"')
    for run in para.runs:
        clr = None
        try:
            clr = run.font.color.rgb
        except:
            pass
        print(f'  Run: bold={run.bold}, size={run.font.size}, name={run.font.name}, color={clr}')

print()
print(f'=== TABLES ({len(doc.tables)}) ===')
for ti, table in enumerate(doc.tables):
    print(f'Table {ti}: rows={len(table.rows)}, cols={len(table.columns)}')
    tblPr_xml = table._tbl.tblPr
    
    # Table width
    tblW = tblPr_xml.find(qn('w:tblW'))
    if tblW is not None:
        print(f'  Width: {tblW.get(qn("w:w"))}, type={tblW.get(qn("w:type"))}')
    
    # Table indent
    tblInd = tblPr_xml.find(qn('w:tblInd'))
    if tblInd is not None:
        print(f'  Indent: {tblInd.get(qn("w:w"))}')
    
    # Table style
    tblStyle = tblPr_xml.find(qn('w:tblStyle'))
    if tblStyle is not None:
        print(f'  Style: {tblStyle.get(qn("w:val"))}')
    
    # Print borders raw XML
    tblBorders = tblPr_xml.find(qn('w:tblBorders'))
    if tblBorders is not None:
        print('  Borders:')
        for b in tblBorders:
            tag = b.tag.split('}')[-1]
            val = b.get(qn('w:val'))
            sz = b.get(qn('w:sz'))
            color = b.get(qn('w:color'))
            print(f'    {tag}: val={val}, sz={sz}, color={color}')
    
    # Column widths
    tblGrid = table._tbl.tblGrid
    if tblGrid is not None:
        cols = [c.get(qn('w:w')) for c in tblGrid.findall(qn('w:gridCol'))]
        print(f'  Column widths (twips): {cols}')
    
    # Row heights and cell details
    for ri, row in enumerate(table.rows):
        trPr = row._tr.trPr
        trHeight = trPr.find(qn('w:trHeight')) if trPr is not None else None
        ht = trHeight.get(qn('w:val')) if trHeight is not None else 'auto'
        header_flag = trPr.find(qn('w:tblHeader')) if trPr is not None else None
        is_header = header_flag is not None
        print(f'  Row {ri}: height={ht}, isHeader={is_header}')
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.tcPr
            shd = tcPr.find(qn('w:shd')) if tcPr is not None else None
            bg = shd.get(qn('w:fill')) if shd is not None else 'none'
            
            # vMerge
            vMerge = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
            vmerge_val = vMerge.get(qn('w:val')) if vMerge is not None else None
            
            # Cell width
            tcW = tcPr.find(qn('w:tcW')) if tcPr is not None else None
            cw = tcW.get(qn('w:w')) if tcW is not None else 'auto'
            
            # Margins
            tcMar = tcPr.find(qn('w:tcMar')) if tcPr is not None else None
            
            txt = cell.text[:40].replace('\n', '|')
            print(f'    Cell ({ri},{ci}): bg={bg}, width={cw}, vMerge={vmerge_val}, text="{txt}"')
            for para in cell.paragraphs[:1]:
                pPr = para._p.pPr
                jc = pPr.find(qn('w:jc')) if pPr is not None else None
                align = jc.get(qn('w:val')) if jc is not None else 'left'
                pBdr = pPr.find(qn('w:pBdr')) if pPr is not None else None
                sp_bef = para.paragraph_format.space_before
                sp_aft = para.paragraph_format.space_after
                line_sp = para.paragraph_format.line_spacing
                print(f'      Para: align={align}, sp_before={sp_bef}, sp_after={sp_aft}, line_sp={line_sp}')
                for run in para.runs[:1]:
                    clr = None
                    try: clr = run.font.color.rgb
                    except: pass
                    print(f'      Run: bold={run.bold}, italic={run.italic}, size={run.font.size}, name={run.font.name}, color={clr}')
