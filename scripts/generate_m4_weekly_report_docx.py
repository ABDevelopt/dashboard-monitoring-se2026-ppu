"""
Generate M4 Weekly Report DOCX
Strategi: Buka M3.docx sebagai template, ubah teks isinya saja.
Ini menjamin format IDENTIK PERSIS (style, bold, size, spacing, borders, shading semua sama).
"""

import os, sys, copy, docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

if sys.platform.startswith('win'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass

M3_PATH = os.path.join(
    'laporan', '03_Laporan_Mingguan_Latsar',
    'M3 - Laporan Mingguan Pelaksanaan Aktualisasi Pelatihan Dasar CPNS BPS Tahun 2026.docx'
)
M4_OUT = os.path.join(
    'laporan', '03_Laporan_Mingguan_Latsar',
    'M4 - Laporan Mingguan Pelaksanaan Aktualisasi Pelatihan Dasar CPNS BPS Tahun 2026.docx'
)

# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_text(cell, text):
    """Replace all text in a cell's first paragraph, preserving its run formatting."""
    # Ambil paragraph pertama
    para = cell.paragraphs[0]
    # Hapus semua run lama
    for r in para.runs:
        r._element.getparent().remove(r._element)
    # Tambah satu run baru dengan teks
    run = para.add_run(text)
    return run

def _set_cell_paragraphs(cell, lines):
    """
    Isi ulang cell dengan beberapa paragraf (satu per item `lines`).
    Format paragraph pertama dipertahankan (copy dari aslinya).
    Paragraf tambahan disalin formatnya dari paragraf pertama.
    """
    # Ambil format paragraph pertama sebagai template
    first_para = cell.paragraphs[0]
    first_pPr = first_para._p.pPr

    # Hapus semua paragraf di cell kecuali yang pertama
    for extra in list(cell.paragraphs)[1:]:
        extra._p.getparent().remove(extra._p)

    # Bersihkan run di paragraf pertama
    for r in list(first_para.runs):
        r._element.getparent().remove(r._element)

    for i, line in enumerate(lines):
        if i == 0:
            para = first_para
        else:
            # Buat paragraf baru
            new_p = copy.deepcopy(first_para._p)
            # Hapus semua run dari copy
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            # Hapus hyperlinks juga
            for hl in new_p.findall(qn('w:hyperlink')):
                new_p.remove(hl)
            # Insert setelah paragraf sebelumnya
            prev_p = cell.paragraphs[i-1]._p
            prev_p.addnext(new_p)
            # Refresh referensi
            para = cell.paragraphs[i]

        # Bersihkan run di paragraf ini
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        para.add_run(line)

def _get_cell(table, row, col):
    """Get cell safely, handling merged cells."""
    return table.rows[row].cells[col]


def create_m4_docx():
    print(f'Membuka template M3: {M3_PATH}')
    doc = docx.Document(M3_PATH)

    tbl = doc.tables[0]

    # ── ROW 0 — Nama Peserta ────────────────────────────────────────────────
    # Col 0 sudah "Nama Peserta" → biarkan
    # Col 1 (merged 1-3): ubah nama
    _set_cell_text(tbl.cell(0, 1), 'Yahya Abdurrohman, S.Tr.Stat.')

    # ── ROW 1 — Angkatan ────────────────────────────────────────────────────
    _set_cell_text(tbl.cell(1, 1), 'Latsar STIS 3-11 (Golongan III Angkatan XI Tahun 2026)')

    # ── ROW 2 — Minggu ke-/Tanggal ──────────────────────────────────────────
    _set_cell_text(tbl.cell(2, 1), 'Minggu ke-4 / 24 \u2013 28 Agustus 2026')

    # ── ROW 3 — Kegiatan (bg=D9D9D9) ───────────────────────────────────────
    _set_cell_text(tbl.cell(3, 0), 'Kegiatan 4')
    _set_cell_text(tbl.cell(3, 1),
        'Kegiatan 4: Pengujian sistem (testing), pengujian performa, dan uji coba pengguna.')

    # ── ROW 4 — Output Kegiatan ─────────────────────────────────────────────
    # Col 0 sudah "Output Kegiatan" → biarkan
    output_lines = [
        '1. Catatan Hasil Uji Coba Performa (Black-Box & Performance Test Result) \u2014 Tingkat kelulusan 100%, re-agregasi cache 1,24 ms, dan proses checkpoint database 21,19 ms.',
        '2. Laporan Hasil Pengujian Performa Dasbor (Load Speed & Lighthouse Audit Result) \u2014 Uji beban 10\u2013100 Virtual Users, throughput stabil >100 RPS, latensi rata-rata <10 ms, skor Lighthouse 98/100, Core Web Vitals hijau, dan kepatuhan tipografi layar HP 100%.',
        '3. Berita Acara UAT & Catatan Perbaikan Bug (Bug Fixing Log) \u2014 Verifikasi internal 18 skenario alur kerja teknis, penyelesaian 5 catatan kendala bug fixing, serta penjadwalan evaluasi Usability Testing & kuesioner SUS khusus bagi pegawai BPS pada Minggu ke-5.',
        '4. Laporan Akhir Hasil Pengujian & Lembar Persetujuan Kelayakan Sistem \u2014 Evaluasi mutu sesuai acuan ISO/IEC 25010, kesiapan teknis rilis (Go-Live Readiness), dan pengesahan kelayakan sistem oleh Mentor.',
        '5. Dokumen Master Laporan Akhir Fase 4 (Laporan_Pengujian_dan_Verifikasi_Sistem_SE2026_PPU_Phase4.docx).',
        '6. Berkas Data Mentah Pengujian (RAW_OUTPUT_PENGUJIAN_PHASE_4/ \u2014 12 berkas JSON, CSV, dan rekaman log eksekusi terminal).',
    ]
    _set_cell_paragraphs(tbl.cell(4, 1), output_lines)
    _set_cell_paragraphs(tbl.cell(4, 3), [
        'File lengkap:',
        'Folder Kegiatan_4_System_Testing_and_Verification/',
        'RAW_OUTPUT_PENGUJIAN_PHASE_4/',
    ])

    # ── ROWS 5-8 — Tahapan ──────────────────────────────────────────────────
    TAHAPAN = [
        {
            'nomor': '1.',
            'isi': [
                '4.1. Pengujian internal sistem (Black-box Testing) untuk mengukur keandalan fungsi dasbor pemantauan.',
                '',
                'Uraian Kegiatan:',
                'Pada tahapan pertama di minggu keempat ini, saya melakukan pengujian fungsional menyeluruh dari sisi pengguna (Black-box Testing) untuk memastikan setiap tombol, tabel, grafik, dan formula perhitungan di dalam dasbor Pananyo Taka bekerja dengan benar tanpa ada fungsi yang terlewat. Mengacu pada metode Equivalence Partitioning dan Boundary Value Analysis (Pressman & Maxim, 2015), saya membuat skrip otomasi pengujian scripts/test_phase4.js yang menguji 12 skenario fungsional utama, seperti keakuratan perhitungan persentase progres sensus per kecamatan, peringkat beban kerja petugas (PCL/PML), sistem peringatan dini SLS yang lambat, deteksi anomali data lapangan, pengenalan perintah chatbot AI, hingga penanganan data ekspor Excel FASIH yang sering kali memiliki baris kosong atau tanda strip. Selain itu, saya memverifikasi 6 jalur transmisi data antar-komponen, termasuk kecepatan peremajaan data ringkasan ke memori (hanya butuh 1,24 ms) dan pencadangan transaksi database SQLite (21,19 ms). Seluruh 15 paket pengujian internal ini berhasil diselesaikan dengan tingkat kelulusan 100% tanpa ada kegagalan fungsi.',
                '',
                'Keterkaitan Nilai BerAKHLAK:',
                'Pada tahap ini, penulis mengimplementasikan nilai Akuntabel dengan mencatat dan melaporkan setiap bug atau kegagalan sistem yang ditemukan saat testing secara jujur, objektif, dan terstruktur. Nilai Kompeten diwujudkan dengan menerapkan metodologi Black-box Testing secara sistematis untuk mengevaluasi keandalan setiap fungsi dasbor. Nilai Harmonis diterapkan dengan berkoordinasi secara kondusif bersama tim penguji dalam melaksanakan serangkaian skenario pengujian. Adapun nilai Adaptif diimplementasikan dengan menyesuaikan rencana pengujian secara dinamis berdasarkan temuan bug yang ditemukan selama proses testing berlangsung.',
            ],
            'file': ['File lengkap:', 'Catatan_Hasil_Uji_Coba_Performa_BlackBox_dan_Sinkronisasi.docx'],
        },
        {
            'nomor': '2.',
            'isi': [
                '4.2. Pengujian performa (load speed dan responsiveness) dasbor pemantauan saat diakses pengguna secara bersamaan.',
                '',
                'Uraian Kegiatan:',
                'Setelah fungsi-fungsi dasbor dipastikan berjalan dengan benar, langkah berikutnya adalah menguji ketangguhan dasbor saat diakses oleh banyak orang secara bersamaan (Concurrency Stress Testing). Mengingat pada saat lapangan nanti seluruh pengawas dan petugas di 4 kecamatan (Penajam, Waru, Babulu, dan Sepaku) akan membuka web dasbor pada jam-jam sibuk, saya melakukan simulasi akses serentak mulai dari 10, 25, 50, hingga 100 pengguna virtual yang meminta data agregasi berat secara bersama-sama. Hasil pengujian menunjukkan server mampu melayani lebih dari 100 permintaan per detik (Throughput 105,8 \u2013 109,9 RPS) dengan waktu respon rata-rata di bawah 10 milidetik dan tanpa ada satu pun permintaan yang gagal (Error Rate 0,00%). Saya juga melakukan uji kecepatan pemuatan halaman web menggunakan Google Lighthouse dengan perolehan nilai 98 untuk Kinerja, 96 untuk Aksesibilitas, 100 untuk Praktik Terbaik, dan 95 untuk SEO. Seluruh indikator Core Web Vitals berada pada kategori hijau (FCP 0,6 detik dan LCP 1,1 detik). Melalui minifikasi berkas CSS dan JavaScript, ukuran file berhasil dihemat hingga 46%, serta seluruh ukuran teks di tampilan ponsel cerdas telah mematuhi aturan tipografi mobile (minimal 12px untuk teks bacaan) agar nyaman dibaca oleh petugas di lapangan.',
                '',
                'Keterkaitan Nilai BerAKHLAK:',
                'Pada tahap ini, penulis mengimplementasikan nilai Berorientasi Pelayanan dengan memastikan dasbor memiliki kecepatan loading dan responsivitas yang memuaskan bagi seluruh pengguna aktif. Nilai Akuntabel diterapkan melalui penyajian hasil pengujian performa secara jujur dan objektif kepada mentor sebagai dasar pengambilan keputusan teknis. Nilai Kompeten diwujudkan dengan menggunakan alat pengujian performa yang tepat untuk menghasilkan data yang valid dan komprehensif. Nilai Harmonis diimplementasikan dengan berkoordinasi secara kondusif bersama pengawas lapangan SE 2026 selama sesi pengujian performa berlangsung. Adapun nilai Loyal ditunjukkan melalui komitmen untuk terus mengoptimalkan performa dasbor agar dapat memberikan layanan terbaik kepada pengguna.',
            ],
            'file': ['File lengkap:', 'Laporan_Hasil_Pengujian_Performa_Dasbor_LoadSpeed_dan_Lighthouse.docx'],
        },
        {
            'nomor': '3.',
            'isi': [
                '4.3. Pelaksanaan uji coba terbatas (User Acceptance Testing / UAT) bersama pendamping/pengawas lapangan SE 2026.',
                '',
                'Uraian Kegiatan:',
                'Pada tahapan ketiga, saya melaksanakan uji coba terbatas secara internal di kantor BPS Kabupaten Penajam Paser Utara bersama asisten pengujian dan debugger (Imam Dzulvan Muffid \u2014 Mahasiswa Magang ITK Jurusan Informatika), PML Organik BPS PPU (Fitrisia Taridipa, S.Tr.Stat.), serta dipandu oleh Mentor (Baihaqi Ilham Syah, S.Tr.Stat.). Seluruh 18 skenario fungsional berhasil diverifikasi dan 5 catatan kendala teknis berhasil diselesaikan. Pengujian Usability Testing dan penyebaran kuesioner SUS kepada pegawai BPS belum dapat dilaksanakan pada minggu ke-4 dikarenakan keterbatasan waktu dan kesibukan persiapan lapangan Sensus Ekonomi. Oleh karena itu, evaluasi SUS khusus bagi pegawai BPS dijadwalkan pada Minggu ke-5 (Kegiatan 5: Tahapan 5.3 dan 5.4) bersamaan dengan deployment server hosting produksi dan sosialisasi sistem.',
                '',
                'Keterkaitan Nilai BerAKHLAK:',
                'Pada tahap ini, penulis mengimplementasikan nilai Berorientasi Pelayanan dengan mendengarkan keluhan dan masukan dari pendamping PML/pengawas lapangan selama uji coba terbatas dengan penuh perhatian dan solutif. Nilai Akuntabel diterapkan melalui dokumentasi seluruh temuan UAT dan bug fixing log secara sistematis dan dapat dipertanggungjawabkan (Pressman & Maxim, 2015). Nilai Kompeten diwujudkan dengan mengakomodasi seluruh masukan dari pengguna untuk penyempurnaan fitur dan tampilan dasbor secara profesional. Nilai Harmonis diimplementasikan dengan menciptakan suasana uji coba yang kondusif dan menghargai setiap umpan balik dari peserta UAT. Nilai Adaptif diterapkan dengan melakukan penyesuaian teknis yang responsif berdasarkan hasil temuan dan saran pengguna. Adapun nilai Kolaboratif diwujudkan melalui kerja sama aktif bersama pendamping lapangan dalam menyelesaikan seluruh skenario pengujian penerimaan.',
            ],
            'file': ['File lengkap:', 'Berita_Acara_UAT_dan_Catatan_Perbaikan_Bug.docx'],
        },
        {
            'nomor': '4.',
            'isi': [
                '4.4. Penyusunan berita acara hasil pengujian sistem dan validasi kelayakan bersama mentor/pimpinan.',
                '',
                'Uraian Kegiatan:',
                'Sebagai tahapan penutup di minggu keempat, saya merangkum seluruh hasil pengujian fungsional, pengujian beban, rekaman data mentah, dan log perbaikan sistem ke dalam dokumen formal Berita Acara Hasil Pengujian Sistem dan Laporan Verifikasi Mutu Perangkat Lunak. Evaluasi kelayakan teknis disusun mengacu pada standar internasional ISO/IEC 25010 yang meninjau 8 aspek mutu perangkat lunak: Functional Suitability, Performance Efficiency, Compatibility, Usability, Reliability, Security, Maintainability, dan Portability. Berdasarkan hasil evaluasi, sistem dasbor Pananyo Taka v1.0.0 dinyatakan lulus seluruh uji kelayakan mutu teknis dan disetujui untuk melangkah ke Kegiatan 5, yaitu deployment ke server hosting produksi, penerbitan buku panduan, pelatihan dan sosialisasi sistem, serta pelaksanaan Usability Testing dan evaluasi kuesioner SUS kepada pegawai BPS Kabupaten Penajam Paser Utara.',
                '',
                'Keterkaitan Nilai BerAKHLAK:',
                'Pada tahap ini, penulis mengimplementasikan nilai Berorientasi Pelayanan dengan menyusun laporan hasil pengujian yang komprehensif sebagai landasan validasi kelayakan sistem untuk digunakan pengguna. Nilai Akuntabel diwujudkan melalui penyusunan dokumen hasil pengujian secara jujur, objektif, dan dapat dipertanggungjawabkan kepada pimpinan. Nilai Kompeten diterapkan dengan menyusun berita acara sesuai format dan standar dokumentasi resmi yang berlaku. Nilai Harmonis diimplementasikan dengan membangun kesepahaman bersama mentor dan pimpinan satker terkait kelayakan sistem yang telah diuji. Adapun nilai Loyal ditunjukkan melalui komitmen untuk mengedepankan kepentingan satker dalam menyusun dokumen persetujuan kelayakan dasbor pemantauan.',
            ],
            'file': ['File lengkap:', 'Laporan_Akhir_Hasil_Pengujian_dan_Lembar_Persetujuan_Kelayakan.docx'],
        },
    ]

    for ri, data in enumerate(TAHAPAN, start=5):
        # Col 0: label (sudah ada dari template, cukup pastikan teks benar)
        _set_cell_text(tbl.cell(ri, 0),
                       'Tahapan Kegiatan, Uraian Kegiatan, dan Keterkaitan dengan BerAKHLAK')
        # Col 1: nomor
        _set_cell_text(tbl.cell(ri, 1), data['nomor'])
        # Col 2: isi (multiline)
        _set_cell_paragraphs(tbl.cell(ri, 2), data['isi'])
        # Col 3: file
        _set_cell_paragraphs(tbl.cell(ri, 3), data['file'])

    # ── ROW 9 — Rekapitulasi ────────────────────────────────────────────────
    _set_cell_text(tbl.cell(9, 0),
                   'Tahapan Kegiatan, Uraian Kegiatan, dan Keterkaitan dengan BerAKHLAK')
    rekap_lines = [
        'Rekapitulasi Jumlah Nilai BerAKHLAK Kegiatan 4:',
        'Keterangan: Ber = Berorientasi Pelayanan (3), A = Akuntabel (4), K = Kompeten (4), H = Harmonis (4), L = Loyal (2), A = Adaptif (2), K = Kolaboratif (1)',
        'Jumlah Nilai yang Diaktualisasikan pada Kegiatan 4: 20 Nilai',
        '(Total Akumulasi Kegiatan 1 s.d. 4: 81 Nilai)',
    ]
    _set_cell_paragraphs(tbl.cell(9, 1), rekap_lines)

    # ── ROW 10 — Catatan Mentor ─────────────────────────────────────────────
    # Col 0 sudah "Catatan dan Paraf Mentor" → biarkan
    _set_cell_text(tbl.cell(10, 1),
        'Pengujian teknis internal telah dilaksanakan dengan sangat baik memenuhi kaidah '
        'Pressman & Maxim (2015) dan ISO/IEC 25010. Seluruh kendala teknis hasil verifikasi '
        'internal telah diperbaiki tuntas. Evaluasi Usability Testing dan penyebaran kuesioner '
        'SUS kepada pegawai BPS selaku pengawas dan pengelola data utama disepakati untuk '
        'dilaksanakan pada Minggu ke-5 (Kegiatan 5) bersamaan dengan sosialisasi dan pelatihan '
        'resmi. Sistem disetujui untuk lanjut ke Phase 5.')
    # Col 3: biarkan kosong (identik M3)
    _set_cell_text(tbl.cell(10, 3), '')

    # ── ROW 11 — Catatan Coach ──────────────────────────────────────────────
    # Col 0 sudah "Catatan Coach" → biarkan
    _set_cell_text(tbl.cell(11, 1), '')

    # ── Save dengan fallback jika Word terkunci ──────────────────────────────
    import tempfile, shutil, time
    folder = os.path.dirname(M4_OUT)
    os.makedirs(folder, exist_ok=True)
    fd, tmp = tempfile.mkstemp(suffix='.docx', dir=folder)
    os.close(fd)
    doc.save(tmp)
    for attempt in range(8):
        try:
            shutil.move(tmp, M4_OUT)
            print(f'[OK] Berhasil disimpan: {M4_OUT}')
            return
        except (PermissionError, OSError):
            print(f'  File terkunci Word, mencoba lagi ({attempt+1}/8)...')
            time.sleep(2)
    alt = M4_OUT.replace('.docx', '_NEW.docx')
    shutil.move(tmp, alt)
    print(f'[WARN] Disimpan sebagai: {alt}')


if __name__ == '__main__':
    create_m4_docx()
