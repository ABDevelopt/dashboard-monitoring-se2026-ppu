"""
Verifikasi bahwa format M4 identik dengan M3 (non-Updated).
Membandingkan struktur tabel: style, col widths, row heights, bg, vMerge, font size, bold.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

def inspect(path, label):
    doc = Document(path)
    print(f'\n{"="*60}')
    print(f'  {label}')
    print(f'{"="*60}')
    s = doc.sections[0]
    print(f'  Margin: top={s.top_margin} left={s.left_margin}')
    print(f'  Page: {s.page_width} x {s.page_height}')

    p0 = doc.paragraphs[0]
    r0 = p0.runs[0]
    print(f'  Title: align={p0.alignment}, bold={r0.bold}, size={r0.font.size}, font={r0.font.name}')

    t = doc.tables[0]
    tblStyle = t._tbl.tblPr.find(qn('w:tblStyle'))
    print(f'  Table style: {tblStyle.get(qn("w:val")) if tblStyle is not None else "none"}')
    grid = [c.get(qn('w:w')) for c in t._tbl.tblGrid.findall(qn('w:gridCol'))]
    print(f'  Col widths: {grid}')

    for ri in range(12):
        row = t.rows[ri]
        trH = row._tr.trPr.find(qn('w:trHeight')) if row._tr.trPr is not None else None
        ht = trH.get(qn('w:val')) if trH is not None else '-'
        cells_info = []
        for ci in range(4):
            if ci >= len(row.cells): break
            cell = row.cells[ci]
            tcPr = cell._tc.tcPr
            shd = tcPr.find(qn('w:shd')) if tcPr is not None else None
            bg = shd.get(qn('w:fill')) if shd is not None else '-'
            vm = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
            vm_val = (vm.get(qn('w:val')) or 'cont') if vm is not None else '-'
            txt = cell.text[:20].replace('\n','|')
            r = cell.paragraphs[0].runs[0] if cell.paragraphs and cell.paragraphs[0].runs else None
            bold = r.bold if r else '-'
            size = r.font.size if r else '-'
            cells_info.append(f'C{ci}[bg={bg} vm={vm_val} bold={bold} sz={size} "{txt}"]')
        print(f'  R{ri:02d} h={ht}: {" | ".join(cells_info)}')

BASE = r'laporan\03_Laporan_Mingguan_Latsar'
inspect(
    rf'{BASE}\M3 - Laporan Mingguan Pelaksanaan Aktualisasi Pelatihan Dasar CPNS BPS Tahun 2026.docx',
    'M3 (ACUAN)'
)
inspect(
    rf'{BASE}\M4 - Laporan Mingguan Pelaksanaan Aktualisasi Pelatihan Dasar CPNS BPS Tahun 2026.docx',
    'M4 (HASIL)'
)
